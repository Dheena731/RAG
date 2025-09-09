'''Document processor for handling various file types (PDF, DOCX, TXT, images)'''
import io
import base64
from pathlib import Path
from typing import List, Dict, Any, Union
import logging

import PyPDF2
from PIL import Image
from docx import Document
import pytesseract  # For OCR on images

logger = logging.getLogger(__name__)


class DocumentProcessor:
    def __init__(self, config):
        self.config = config
        self.chunk_size = config.CHUNK_SIZE
        self.chunk_overlap = config.CHUNK_OVERLAP
    
    def detect_file_type(self, file_path: Path) -> str:
        """Detect the type of file based on extension"""
        suffix = file_path.suffix.lower()
        
        if suffix in self.config.SUPPORTED_TEXT_FORMATS:
            if suffix == ".pdf":
                return "pdf"
            elif suffix == ".docx":
                return "docx"
            elif suffix in [".txt", ".md"]:
                return "text"
        elif suffix in self.config.SUPPORTED_IMAGE_FORMATS:
            return "image"
        
        raise ValueError(f"Unsupported file type: {suffix}")
    
    def process_file(self, file_path: Path, file_type: str) -> List[Dict[str, Any]]:
        """Process a file and return chunks of content"""
        logger.info(f"Processing {file_type} file: {file_path}")
        
        if file_type == "pdf":
            return self._process_pdf(file_path)
        elif file_type == "docx":
            return self._process_docx(file_path)
        elif file_type == "text":
            return self._process_text(file_path)
        elif file_type == "image":
            return self._process_image(file_path)
        else:
            raise ValueError(f"Unknown file type: {file_type}")
    
    def _process_pdf(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extract text from PDF and create chunks"""
        chunks = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        # Split text into chunks
                        text_chunks = self._split_text(text)
                        
                        for chunk_idx, chunk_text in enumerate(text_chunks):
                            chunks.append({
                                "content": chunk_text,
                                "metadata": {
                                    "source": str(file_path),
                                    "page": page_num + 1,
                                    "chunk": chunk_idx,
                                    "file_type": "pdf",
                                    "total_pages": len(pdf_reader.pages)
                                },
                                "type": "text"
                            })
        
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            raise
        
        return chunks
    
    def _process_docx(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extract text from DOCX and create chunks"""
        chunks = []
        
        try:
            doc = Document(file_path)
            full_text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Join all paragraphs
            text = "\\n".join(full_text)
            
            if text.strip():
                # Split text into chunks
                text_chunks = self._split_text(text)
                
                for chunk_idx, chunk_text in enumerate(text_chunks):
                    chunks.append({
                        "content": chunk_text,
                        "metadata": {
                            "source": str(file_path),
                            "chunk": chunk_idx,
                            "file_type": "docx",
                            "total_paragraphs": len(doc.paragraphs)
                        },
                        "type": "text"
                    })
        
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            raise
        
        return chunks
    
    def _process_text(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process plain text files"""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if text.strip():
                # Split text into chunks
                text_chunks = self._split_text(text)
                
                for chunk_idx, chunk_text in enumerate(text_chunks):
                    chunks.append({
                        "content": chunk_text,
                        "metadata": {
                            "source": str(file_path),
                            "chunk": chunk_idx,
                            "file_type": file_path.suffix[1:],  # Remove the dot
                            "total_chars": len(text)
                        },
                        "type": "text"
                    })
        
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
            raise
        
        return chunks
    
    def _process_image(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process image files - convert to base64 and extract text via OCR"""
        chunks = []
        
        try:
            # Load and encode image
            with Image.open(file_path) as img:
                # Convert to RGB if necessary
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Convert to base64 for storage
                buffer = io.BytesIO()
                img.save(buffer, format='JPEG')
                img_base64 = base64.b64encode(buffer.getvalue()).decode()
                
                # Try to extract text using OCR
                try:
                    ocr_text = pytesseract.image_to_string(img)
                    ocr_text = ocr_text.strip()
                except Exception as ocr_error:
                    logger.warning(f"OCR failed for {file_path}: {str(ocr_error)}")
                    ocr_text = ""
                
                # Create image chunk
                chunks.append({
                    "content": img_base64,  # Store image as base64
                    "metadata": {
                        "source": str(file_path),
                        "file_type": "image",
                        "format": img.format or file_path.suffix[1:],
                        "size": img.size,
                        "mode": img.mode,
                        "ocr_text": ocr_text  # Store extracted text as metadata
                    },
                    "type": "image"
                })
                
                # If OCR extracted meaningful text, create a separate text chunk
                if len(ocr_text) > 20:  # Only if substantial text was found
                    text_chunks = self._split_text(ocr_text)
                    for chunk_idx, chunk_text in enumerate(text_chunks):
                        chunks.append({
                            "content": chunk_text,
                            "metadata": {
                                "source": str(file_path),
                                "chunk": chunk_idx,
                                "file_type": "image_text",
                                "extracted_from": "OCR"
                            },
                            "type": "text"
                        })
        
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {str(e)}")
            raise
        
        return chunks
    
    def _split_text(self, text: str) -> List[str]:
        """Split text into chunks with overlap"""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Find the end position
            end = start + self.chunk_size
            
            if end >= len(text):
                # Last chunk
                chunks.append(text[start:])
                break
            
            # Try to find a good breaking point (sentence ending or paragraph)
            break_point = end
            
            # Look for paragraph breaks first
            paragraph_break = text.rfind('\\n\\n', start, end)
            if paragraph_break > start:
                break_point = paragraph_break + 2
            else:
                # Look for sentence endings
                sentence_endings = ['. ', '! ', '? ']
                best_break = -1
                
                for ending in sentence_endings:
                    pos = text.rfind(ending, start, end)
                    if pos > best_break:
                        best_break = pos + len(ending)
                
                if best_break > start:
                    break_point = best_break
            
            chunks.append(text[start:break_point])
            
            # Move start position with overlap
            start = break_point - self.chunk_overlap
            if start < 0:
                start = break_point
        
        return chunks
